"""从核心 docx 文件中抽取纯文本，便于检索关键证据。"""
import os
import sys
from pathlib import Path
from docx import Document

BASE = Path("/home/user/workspace/P03_SAR_AIS_DarkVessel_RemoteSensing/02_支撑材料包/RS_SAR_AIS_DarkVessel_Export_20260517/01_Core_Source_Documents")
OUT = Path("/home/user/workspace/paper_v2/extracted_text")
OUT.mkdir(parents=True, exist_ok=True)

# 优先级最高的 4 份文档
PRIORITY = [
    "13-海上搜救与应急指挥科研平台-软件设计说明书.docx",       # Methods 核心
    "2-海上搜救与应急指挥科研平台-软件测评报告.docx",            # Validation 核心
    "11-海上搜救与应急指挥科研平台-数据说明书.docx",            # Data 核心
    "16-海上搜救与应急指挥科研平台-技术参数.docx",              # 系统指标
    "14-海上搜救与应急指挥科研平台-实施方案.docx",              # 总体方案
    "1-海上搜救与应急指挥科研平台-空间遥感数据清单.docx",        # 数据清单
    "1-海上搜救与应急指挥科研平台-软件系统测试大纲.docx",       # 测试大纲
    "1-海上搜救与应急指挥科研平台-软件系统工作总结报告.docx",   # 工作总结
]

def extract(path):
    doc = Document(path)
    lines = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if t:
            lines.append(t)
    # 提取表格
    for ti, tbl in enumerate(doc.tables):
        lines.append(f"\n[TABLE {ti}]")
        for row in tbl.rows:
            cells = [c.text.strip().replace("\n", " ") for c in row.cells]
            lines.append(" | ".join(cells))
        lines.append(f"[/TABLE {ti}]\n")
    return "\n".join(lines)

for fname in PRIORITY:
    src = BASE / fname
    if not src.exists():
        print(f"MISSING: {fname}")
        continue
    out_name = fname.replace(".docx", ".txt")
    out_path = OUT / out_name
    print(f"Extracting {fname} ({src.stat().st_size/1024/1024:.1f} MB)...")
    try:
        txt = extract(src)
        out_path.write_text(txt, encoding="utf-8")
        print(f"  -> {len(txt)} chars to {out_path.name}")
    except Exception as e:
        print(f"  ERROR: {e}")
